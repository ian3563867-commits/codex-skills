#!/usr/bin/env python3
"""Create an eCATCH meeting-minutes DOCX from the bundled template."""

from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = (
    SKILL_DIR
    / "assets"
    / "templates"
    / "0151-世祥斗南-MML-PMS、出貨流程確認-20210917.docx"
)
DEFAULT_OUTPUT_DIR = Path.cwd()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE))
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--output-name", required=True)
    parser.add_argument("--meeting-name", required=True)
    parser.add_argument("--place", required=True)
    parser.add_argument("--time", required=True)
    parser.add_argument("--chair", required=True)
    parser.add_argument("--recorder", default="Ian")
    parser.add_argument("--attendees", default="")
    parser.add_argument("--content", required=True)
    return parser.parse_args()


def clear_keep_cell_properties(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)
    return cell.add_paragraph()


def apply_run_font(run, size: float = 10.5) -> None:
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)


def set_text(cell, text: str, size: float = 10.5) -> None:
    p = clear_keep_cell_properties(cell)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if text:
        run = p.add_run(text)
        apply_run_font(run, size)


def set_multiline(cell, text: str) -> None:
    p = clear_keep_cell_properties(cell)
    for idx, line in enumerate(text.strip().splitlines()):
        if idx:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(line)
        apply_run_font(run, 11)


def remove_exact_heights_for_editable_rows(table) -> None:
    for row_idx in range(0, min(6, len(table.rows))):
        row = table.rows[row_idx]
        tr_pr = row._tr.trPr
        if tr_pr is not None:
            for height in list(tr_pr.findall(qn("w:trHeight"))):
                tr_pr.remove(height)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for cell in table.rows[5].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def count_template_objects(docx_path: Path) -> tuple[int, int]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    drawings = 0
    picts = 0
    with ZipFile(docx_path) as package:
        for name in package.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                root = etree.fromstring(package.read(name))
                drawings += len(root.xpath(".//w:drawing", namespaces=ns))
                picts += len(root.xpath(".//w:pict", namespaces=ns))
    return drawings, picts


def main() -> None:
    args = parse_args()
    template = Path(args.template)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / args.output_name

    before = count_template_objects(template)
    doc = Document(template)
    table = doc.tables[0]

    set_text(table.cell(0, 1), args.meeting_name)
    set_text(table.cell(1, 1), args.place)
    set_text(table.cell(1, 3), args.time)
    set_text(table.cell(2, 1), args.chair)
    set_text(table.cell(2, 3), args.recorder)
    set_text(table.cell(3, 1), args.attendees)
    set_multiline(table.cell(5, 0), args.content)
    remove_exact_heights_for_editable_rows(table)

    doc.save(output)
    after = count_template_objects(output)

    appendix_text = table.rows[6].cells[0].text if len(table.rows) > 6 else ""
    print(f"output={output}")
    print(f"template_objects_before=drawing:{before[0]},pict:{before[1]}")
    print(f"template_objects_after=drawing:{after[0]},pict:{after[1]}")
    print(f"appendix_present={'附錄' in appendix_text}")

    if after != before:
        raise SystemExit("Header/footer drawing object count changed; inspect before delivery.")
    if "附錄" not in appendix_text:
        raise SystemExit("Appendix row text is missing; inspect before delivery.")


if __name__ == "__main__":
    main()
